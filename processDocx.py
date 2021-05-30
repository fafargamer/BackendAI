import urllib
import spacy
import requests
import string
import docx
import io
import numpy as np
from nltk.tokenize import RegexpTokenizer
from string import digits
from docx import Document   #pip install python-docx
import PyPDF2               #pip install PyPDF2

# import StopWordRemoverFactory class
from Sastrawi.StopWordRemover.StopWordRemoverFactory import StopWordRemoverFactory, StopWordRemover, ArrayDictionary #pip install Sastrawi
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from stopwords import more_stopword as more_St

factory = StopWordRemoverFactory()
# factoryStem = StemmerFactory()
# stopword = factory.create_stop_word_remover()
# stemmer = factoryStem.create_stemmer()
import re # impor modul regular expression

# Ambil Stopword bawaan
stop_factory = StopWordRemoverFactory().get_stop_words()


# Merge stopword
data = stop_factory + more_St

dictionary = ArrayDictionary(data)
stopword = StopWordRemover(dictionary)


def processDocxParagraph(document):
  # global texts
  
  # document=requests.get(inputUrl, allow_redirects=True)
  document=Document(io.BytesIO(document.content))
  document.save('test.docx')
  document = docx.Document('test.docx')
  documentArray=[]
  pureDocumentArray=[]
  for paragraph in document.paragraphs:
    if(len(paragraph.text)==1):
      print(paragraph.text)
    if(len(paragraph.text)>0 and paragraph.text != " "):
      paragraph.text = paragraph.text.replace('\u2013', '-')
      paragraph.text = paragraph.text.replace('\u00a0', '')
      # newText = newText.replace('\u00a0', '')
      
      documentArray.append(paragraph.text)
    # pureDocumentArray.append(paragraph.text)

    #catatan :
    # \u2013 = -
    # \u201c = " (depan)
    # \u201d = " (belakang)


  result = []
  s=""
  for i in range(len(documentArray)):
      text = []
      for j in range(len(documentArray[i])):

        
        tokenizer = RegexpTokenizer(r'\w+')
        textsTokenized=tokenizer.tokenize(documentArray[i][j])
        if(len(textsTokenized) < 1):
          textsTokenized = [' ']

        text.append(textsTokenized[0])
      textor = s.join(text)
      textor = stopword.remove(textor)
      # textor = stemmer.stem(textor)
      textor = re.sub(r"\d+", "", textor)
      result.append(textor)


  dividedDocument = []
  dividedPure = []


  print('Document Processed')


  divisor = 1

  if len(result) > 255:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')

      document_split = np.array_split(result, int(len(result)/(divisor*64)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/(divisor*64)))
      print(document_split[0])
      print(pure_split[0])
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure


  elif len(result) > 127:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')

      document_split = np.array_split(result, int(len(result)/(divisor*32)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/(divisor*32)))
      print(document_split[0])
      print(pure_split[0])
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure


  elif len(result) > 63:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')

      document_split = np.array_split(result, int(len(result)/(divisor*16)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/(divisor*16)))
      print(document_split[0])
      print(pure_split[0])
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure


  elif len(result) > 31:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')

      document_split = np.array_split(result, int(len(result)/(divisor*8)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/(divisor*8)))
      # print(document_split[0])
      # print(pure_split[0])
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure


  elif len(result) > 15:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')

      document_split = np.array_split(result, int(len(result)/(divisor*4)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/(divisor*4)))
      print(document_split[0])
      print(pure_split[0])
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure

  elif len(result) > 7:
    # print(len(result))
    # print('Dokumen dapat dibagi 5')
      document_split = np.array_split(result, int(len(result)/(divisor*2)))
      pure_split = np.array_split(documentArray, int(len(documentArray)/divisor*2))
      # print(document_split)

      for i in range(len(document_split)):
        # print("Part : {}".format(pure_split[i]))
        joinedDocument = ' '.join(document_split[i])
        joinedPure = ' '.join(pure_split[i])
        dividedDocument.append(joinedDocument)
        dividedPure.append(joinedPure)

      result = dividedDocument
      documentArray = dividedPure



  return documentArray, result